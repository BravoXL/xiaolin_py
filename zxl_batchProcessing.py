from docx import Document
from docx.enum.text import  WD_ALIGN_PARAGRAPH #对齐
from docx.shared import Pt,Inches   # 单位
from docx.oxml.ns import qn   # 中文字体
文件 = Document()  # 新建文档
import pandas as pd
# ————————————————————！！！！！！！！！！！！！！
数据 = pd.read_excel('D:/zxl_PreciousPythonFiles/batch.xlsx')
# print(数据.iloc[0,1])
import datetime as dt
日期 = dt.datetime.now()
当前日期 = str(日期.year) + '年' + str(日期.month) + '月' + str(日期.day) + '日' # print(当前日期)
for i in 数据['序号']:  # 笔记8.6
     姓名 = 数据.iloc[i-1,1]
     身份证 = str(数据.iloc[i-1,2])
     电话 = str(数据.iloc[i-1,3])
     住址 = str(数据.iloc[i-1,4])
     来自 = str(数据.iloc[i-1,5])
     段落1 = 文件.add_paragraph()  # 增加一个段落，这个段落是标题
     段落1.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 对齐方式为居中，没有这句话默认左对齐
     标题 = 段落1.add_run('复工证明') # 给段落添加一个块，写上文字
     标题.font.name = "Arial"  # 设置英文字体
     标题.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体') # 设置中文字体
     标题.font.size = Pt(18) # 设置字号
     标题.font.bold = True  # 设置加粗
     段落1.space_after = Pt(5) # 断后距离5磅
     段落1.space_before = Pt(5)  # 断前距离5磅
     段落2 = 文件.add_paragraph()  # 增加第二个段落，这个段落是正文
     正文 = 段落2.add_run(f'兹有大唐天子圣谕，{姓名}同志，身份证号{身份证},联系电话{电话}，系我天朝{住址}子民，根据我天朝复工要求，该同志在 防疫管控期间需要保护唐僧赴西天大雷音寺求取真经，在满足{来自}当地疫情管控要求情况下，请检测体温无异常后予以放行。\n')
     正文.font.name = 'Arial'
     正文.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
     正文.font.size = Pt(14)
     正文.font.bold = True
     段落2.paragraph_format.first_line_indent = Inches(0.4)  # 左缩进0.4英寸
     段落3 = 文件.add_paragraph()  # 增加一个自然段，这段是特此说明
     说明 = 段落3.add_run('特此说明。\n\n\n')
     说明.font.name = 'Arial'
     说明.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
     说明.font.size = Pt(14)
     说明.font.bold = True
     段落3.paragraph_format.first_line_indent = Inches(0.4)
     段落4 = 文件.add_paragraph()  # 增加一个自然段，这段是单位名
     单位 = 段落4.add_run('大唐天子李世民\n')
     单位.font.name = 'Arial'
     单位.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
     单位.font.size = Pt(14)
     单位.font.size = Pt(14)
     单位.font.bold = True
     段落4.paragraph_format.first_line_indent = Inches(4.0)
     段落5 = 文件.add_paragraph()  # 增加一个自然段
     日期 = 段落5.add_run(f'{当前日期}')
     日期.font.name = 'Arial'
     日期.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
     日期.font.size = Pt(14)
     日期.font.bold = True
     段落5.paragraph_format.first_line_indent = Inches(3.9)
     文件.add_page_break()  # 增加分页符
# ————————————————————！！！！！！！！！！！！！！
文件.save(r"D:\zxl_PreciousPythonFiles\zxl_test_Batch.docx")
