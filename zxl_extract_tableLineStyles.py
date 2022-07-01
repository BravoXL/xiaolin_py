#zxl function \\\\*before run, choose the table style by mouse and save,then run.so the code is for extract the style.
from docx.enum.style import WD_STYLE_TYPE
from docx import Document
file1=Document(r"D:\zxl_PreciousPythonFiles\zxl_test2.docx")
styles1=file1.styles
for s in styles1:
    if s.type == WD_STYLE_TYPE.TABLE:
        file1.add_paragraph("tableStyle:"+s.name)
        table1=file1.add_table(3,3,style=s)
        cell1=table1.rows[0].cells
        cell1[0].text="zxl1"
        cell1[1].text = "zxl2"
        cell1[2].text = "zxl3"
        file1.add_paragraph("\n")
file1.save(r"D:\zxl_PreciousPythonFiles\zxl_test2.docx")