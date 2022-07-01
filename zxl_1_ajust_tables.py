 # ****when there is content in cells. the table can't be alignmented.
# \\\\terminal :pip install docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
file1=Document(r"D:\zxl_PreciousPythonFiles\2019摘要合并_06061.docx")
# ____the first column format____
for t in file1.tables:#choose all the tables
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.width = Cm(15.49)

    for r in t.columns[0].cells:#the first column
        r.width=Cm(3.24)
        for k in r.paragraphs:#paragraph align
            k.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
file1.save(r"D:\zxl_PreciousPythonFiles\2019摘要合并_06062.docx")


